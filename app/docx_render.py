# pyright: reportPrivateUsage=false
import re
from dataclasses import dataclass
from datetime import date
from io import BytesIO
from string import ascii_lowercase, ascii_uppercase
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from sqlalchemy.orm import Session

from .list_lines import _ListItem, line_is_list_item, parse_list_for_docx
from .models import Figura, Relatorio, Secao
from .pdf_render import (
    MODELO_DOCX,
    PDF_COVER_IMAGE,
    _RE_FIGURA,
    _RE_TABELA,
    _figura_ids_no_texto,
    _parse_tabela_marker,
    _produto_codigo_capa,
)


TEXT_COLOR = RGBColor(0x14, 0x14, 0x14)
HEADING_COLOR = RGBColor(0x10, 0x22, 0x46)
MUTED_COLOR = RGBColor(0x7A, 0x86, 0x94)


@dataclass(frozen=True)
class _DocxTabelaOpts:
    legenda: str | None = None
    fonte: str | None = None
    posicao: str = "S"


@dataclass(frozen=True)
class _DocxFiguraOpts:
    legenda: str | None = None
    fonte: str | None = None
    posicao: str = "I"


@dataclass
class _MarcadoresContadores:
    fig: int
    tab: int


def _figura_ids_relatorio(rel: Relatorio, section_ids: set[int] | None) -> set[int]:
    ids: set[int] = set()
    for sec in rel.secoes:
        if section_ids is not None and sec.id not in section_ids:
            continue
        for bloco in sec.blocos:
            if bloco.figura_id:
                ids.add(bloco.figura_id)
            ids.update(_figura_ids_no_texto(bloco.conteudo or ""))
    return ids


def _figura_label(sec_numero: str, counter: int) -> str:
    top = (sec_numero or "").split(".")[0]
    return f"{top}.{counter}" if top else str(counter)


def _set_runs_font(
    paragraph,
    *,
    size_pt: float = 10,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor = TEXT_COLOR,
) -> None:
    for run in paragraph.runs:
        run.font.name = "Verdana"
        run.font.size = Pt(size_pt)
        run.font.color.rgb = color
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic


def _format_body_paragraph(paragraph, *, space_after_pt: float = 10) -> None:
    """Parágrafo de corpo. Default 10pt depois para combinar com `margin:0 0 10pt`
    do PDF (`app/templates/pdf/relatorio.html`, regra `.bloco p`)."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(space_after_pt)
    paragraph_format.line_spacing = 1.15
    _set_runs_font(paragraph, size_pt=10)


def _format_list_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Cm(1.27)
    paragraph_format.first_line_indent = Cm(-0.63)
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15
    _set_runs_font(paragraph, size_pt=10)


def _list_extra_indent_cm(level: int) -> float:
    return 0.45 * max(level, 0)


def _int_to_roman_lower(n: int) -> str:
    if n < 1:
        return "i"
    parts = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    sym = "m", "cm", "d", "cd", "c", "xc", "l", "xl", "x", "ix", "v", "iv", "i"
    t, out = n, []
    for v, s in zip(parts, sym):
        while t >= v:
            t -= v
            out.append(s)
    return "".join(out)


def _list_counters_fresh() -> dict[str, int]:
    return {"ol_1": 0, "ol_a": 0, "ol_A": 0, "ol_i": 0, "ol_I": 0}


def _list_prefix(n: _ListItem, ct: dict[str, int]) -> str:
    if n.kind == "ul" or n.kind not in ct:
        return ""
    ct[n.kind] += 1
    c = ct[n.kind]
    k = n.kind
    if k == "ol_1":
        out = f"{c}. "
    elif k == "ol_a":
        ch = ascii_lowercase[(c - 1) % 26] if c <= 26 else "z"
        out = f"{ch}) "
    elif k == "ol_A":
        ch = ascii_uppercase[(c - 1) % 26] if c <= 26 else "Z"
        out = f"{ch}) "
    elif k == "ol_i":
        out = f"{_int_to_roman_lower(c)}. "
    elif k == "ol_I":
        out = f"{_int_to_roman_lower(c).upper()}. "
    else:
        out = ""
    return out


def _docx_add_list_siblings(document: Document, nodes: list[_ListItem], ct: dict[str, int] | None) -> None:
    if not nodes:
        return
    if ct is None:
        ct = _list_counters_fresh()
    for n in nodes:
        ex = _list_extra_indent_cm(n.level)
        if n.kind == "ul":
            p = document.add_paragraph(n.text, style="List Bullet")
            _format_list_paragraph(p)
            p.paragraph_format.left_indent = Cm(1.27 + ex)
            p.paragraph_format.first_line_indent = Cm(-0.63)
        elif n.kind == "ol_1":
            p = document.add_paragraph(n.text, style="List Number")
            _format_list_paragraph(p)
            p.paragraph_format.left_indent = Cm(1.27 + ex)
            p.paragraph_format.first_line_indent = Cm(-0.63)
        else:
            p = document.add_paragraph()
            pre = _list_prefix(n, ct)
            p.add_run(f"{pre}{n.text}")
            _format_body_paragraph(p, space_after_pt=6)
            p.paragraph_format.left_indent = Cm(1.1 + ex)
        if n.children:
            _docx_add_list_siblings(document, n.children, _list_counters_fresh())


def _add_lista_bloco_docx(document: Document, conteudo: str) -> None:
    roots = parse_list_for_docx(conteudo)
    if not roots:
        return
    _docx_add_list_siblings(document, roots, _list_counters_fresh())


def _format_heading(paragraph, level: int) -> None:
    sizes = {1: 13, 2: 12, 3: 11, 4: 10}
    indents = {1: 0, 2: 0, 3: 0.5, 4: 1.0, 5: 1.5, 6: 1.5}
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Cm(indents.get(level, 1.5))
    paragraph_format.space_before = Pt(0 if level == 1 else 10)
    paragraph_format.space_after = Pt(10)
    paragraph_format.line_spacing = 1.15
    paragraph_format.keep_with_next = True
    _set_runs_font(paragraph, size_pt=sizes.get(level, 10), bold=True, italic=level >= 4, color=HEADING_COLOR)


def _add_numbered_heading(document: Document, numero: str, titulo: str, level: int):
    paragraph = document.add_heading("", level=min(level, 4))
    paragraph.add_run(numero or "")
    paragraph.add_run("\t")
    paragraph.add_run((titulo or "").upper() if level == 1 else (titulo or ""))
    _format_heading(paragraph, level)
    indents = {1: 0, 2: 0, 3: 0.5, 4: 1.0, 5: 1.5, 6: 1.5}
    tab_offsets = {1: 0.47, 2: 0.6, 3: 0.73, 4: 0.87, 5: 1.0, 6: 1.0}
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Cm(indents.get(level, 1.5) + tab_offsets.get(level, 3.0))
    )
    return paragraph


def _format_caption(
    paragraph,
    *,
    bold: bool = True,
    space_before_pt: float = 10,
    size_pt: float = 9,
) -> None:
    """Legenda de figura/tabela (e fonte). Defaults agora batem com o PDF:
    9pt bold, 10pt antes e 10pt depois (`.figura .cap`/`.tabela .cap`).
    Caller que precise de variação (ex.: linha 'Fonte:' menor) sobrescreve."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(space_before_pt)
    paragraph_format.space_after = Pt(10)
    paragraph_format.line_spacing = 1.15
    _set_runs_font(paragraph, size_pt=size_pt, bold=bold)


def _format_caption_fonte(paragraph) -> None:
    """Linha 'Fonte: ...' — 8.5pt sem negrito, alinhada ao PDF (`.figura .src`)."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(10)
    paragraph_format.line_spacing = 1.15
    _set_runs_font(paragraph, size_pt=8.5, bold=False)


def _format_empty_section(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(10)
    paragraph_format.line_spacing = 1.15
    _set_runs_font(paragraph, size_pt=10, italic=True, color=MUTED_COLOR)


def _configure_document_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Verdana"
    normal.font.size = Pt(10)
    normal.font.color.rgb = TEXT_COLOR
    normal.paragraph_format.space_after = Pt(12)
    normal.paragraph_format.line_spacing = 1.15

    list_style = styles["List Paragraph"]
    list_style.font.name = "Verdana"
    list_style.font.size = Pt(10)
    list_style.paragraph_format.left_indent = Cm(1.27)
    list_style.paragraph_format.first_line_indent = Cm(-0.63)
    list_style.paragraph_format.space_after = Pt(6)
    list_style.paragraph_format.line_spacing = 1.15

    for level in range(1, 5):
        style = styles[f"Heading {level}"]
        style.font.name = "Verdana"
        style.font.size = Pt({1: 13, 2: 12, 3: 11, 4: 10}[level])
        style.font.bold = True
        style.font.italic = level >= 4
        style.font.color.rgb = HEADING_COLOR
        style.paragraph_format.left_indent = Cm({1: 0, 2: 0, 3: 0.5, 4: 1.0}[level])
        style.paragraph_format.space_before = Pt(0 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(10)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True


def _modelo_asset_bytes(asset_name: str) -> bytes | None:
    try:
        with ZipFile(MODELO_DOCX) as archive:
            return archive.read(f"word/media/{asset_name}")
    except (FileNotFoundError, KeyError):
        return None


def _configure_page(section, *, cover: bool = False) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    if cover:
        section.top_margin = Cm(0)
        section.left_margin = Cm(0)
        section.right_margin = Cm(0)
        section.bottom_margin = Cm(0)
        return
    section.top_margin = Cm(2.22)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.bottom_margin = Cm(2.5)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(1.0)


def _set_paragraph_top_border(paragraph, *, val: str = "dotted", color: str = "CFCFCF", size: str = "6") -> None:
    p_pr = getattr(paragraph, "_p").get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), val)
    top.set(qn("w:sz"), size)
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), color)
    p_bdr.append(top)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_element = getattr(run, "_r")
    run_element.append(fld_char_begin)
    run_element.append(instr_text)
    run_element.append(fld_char_end)


def _configure_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    header.paragraphs[0].text = ""
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logos = _modelo_asset_bytes("image2.png")
    if logos:
        header.paragraphs[0].add_run().add_picture(BytesIO(logos), width=Cm(5.0))

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    _set_paragraph_top_border(paragraph)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(16.8))
    text_run = paragraph.add_run("Plano de Logística e Investimentos do Estado de SP | PLI 2050")
    text_run.font.name = "Verdana"
    text_run.font.size = Pt(7.4)
    text_run.font.color.rgb = RGBColor(0xA8, 0xA8, 0xA8)
    paragraph.add_run("\t")
    _add_page_number(paragraph)
    for run in paragraph.runs[1:]:
        run.font.name = "Verdana"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        run.font.bold = True


def _add_cover(document: Document, rel: Relatorio) -> None:
    _configure_page(document.sections[0], cover=True)
    if PDF_COVER_IMAGE.exists():
        document.add_picture(str(PDF_COVER_IMAGE), width=Cm(21))
    else:
        paragraph = document.add_paragraph("RELATÓRIO MENSAL")
        paragraph.add_run(f"\nProduto {_produto_codigo_capa(rel.codigo)}")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_runs_font(paragraph, size_pt=18, bold=True, color=HEADING_COLOR)
    section = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_page(section)
    _configure_header_footer(section)


def _set_cell_border(cell, *, top: bool = True, bottom: bool = True, color: str = "111111", size: str = "8") -> None:
    tc_pr = getattr(cell, "_tc").get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "bottom"):
        if not ((edge == "top" and top) or (edge == "bottom" and bottom)):
            continue
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        tc_borders.append(element)


def _cell_text(cell, label: str, value: str, *, uppercase_label: bool = True) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.15
    label_run = paragraph.add_run(label.upper() if uppercase_label else label)
    label_run.font.name = "Verdana"
    label_run.font.size = Pt(10.5)
    label_run.font.bold = True
    label_run.font.color.rgb = TEXT_COLOR
    label_run.add_break()
    value_run = paragraph.add_run(value or "")
    value_run.font.name = "Verdana"
    value_run.font.size = Pt(10.5)
    value_run.font.bold = False
    value_run.font.color.rgb = TEXT_COLOR


def _ficha_tecnica_cabecalho(document: Document, rel: Relatorio) -> None:
    heading = document.add_paragraph("Relatório Mensal " + (rel.codigo or ""))
    heading.paragraph_format.space_after = Pt(8.5)
    _set_runs_font(heading, size_pt=15, bold=True, color=TEXT_COLOR)


def _ficha_tecnica_tabela_identificacao(document: Document, rel: Relatorio) -> None:
    rows = [
        ("Código do documento", rel.codigo or ""),
        ("Título", "RELATÓRIO MENSAL"),
        ("Elaboração", "Consórcio Concremat / Transplan"),
        ("Contrato", "Contrato Nº 22.607-5"),
        ("Contratação", "Secretaria de Meio Ambiente, Infraestrutura e Logística - SEMIL - Governo do Estado de São Paulo"),
        ("Financiamento", "Banco Interamericano de Desenvolvimento (BID)"),
        ("Observações", f"Este Relatório corresponde ao entregável {rel.codigo}."),
    ]
    table = document.add_table(rows=len(rows), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, value) in zip(table.rows, rows):
        cell = row.cells[0]
        _set_cell_border(cell, size="11")
        _cell_text(cell, label, value)


def _ficha_tecnica_tabela_versoes(document: Document, rel: Relatorio) -> None:
    versions = document.add_table(rows=4, cols=3)
    versions.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Versão", "Data", "Conteúdo das modificações"]
    values = [["", "", ""], ["", "", ""], [rel.versao or "", date.today().strftime("%d/%m/%Y"), "Versão inicial"]]
    for col_idx, text in enumerate(headers):
        cell = versions.cell(0, col_idx)
        cell.text = text
        _set_cell_border(cell, size="6")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_runs_font(paragraph, size_pt=10.5, bold=True)
    for row_idx, row_values in enumerate(values, start=1):
        for col_idx, text in enumerate(row_values):
            cell = versions.cell(row_idx, col_idx)
            cell.text = text
            _set_cell_border(cell, size="6")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_runs_font(paragraph, size_pt=10.5)


def _add_ficha_tecnica(document: Document, rel: Relatorio) -> None:
    _ficha_tecnica_cabecalho(document, rel)
    _ficha_tecnica_tabela_identificacao(document, rel)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(82)
    _ficha_tecnica_tabela_versoes(document, rel)
    document.add_section(WD_SECTION.NEW_PAGE)
    section = document.sections[-1]
    _configure_page(section)
    _configure_header_footer(section)


def _add_sumario(document: Document, rel: Relatorio, section_ids: set[int] | None) -> None:
    heading = document.add_paragraph("Sumário")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(22.7)
    _set_runs_font(heading, size_pt=13.5, bold=True, color=HEADING_COLOR)
    for sec in rel.secoes:
        if section_ids is not None and sec.id not in section_ids:
            continue
        level = min((sec.numero or "").count(".") + 1, 3)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm({1: 0, 2: 0.388, 3: 0.776}[level])
        paragraph.paragraph_format.space_after = Cm(0.423)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.add_run(sec.numero or "")
        paragraph.add_run("  ")
        paragraph.add_run((sec.titulo or "").upper() if level == 1 else (sec.titulo or ""))
        _set_runs_font(paragraph, size_pt=10.5, italic=(sec.numero or "").count(".") + 1 >= 4)
    document.add_section(WD_SECTION.NEW_PAGE)
    section = document.sections[-1]
    _configure_page(section)
    _configure_header_footer(section)


def _add_assinaturas(document: Document, rel: Relatorio) -> None:
    document.add_section(WD_SECTION.NEW_PAGE)
    section = document.sections[-1]
    _configure_page(section)
    _configure_header_footer(section)
    heading = document.add_paragraph("Página de assinaturas")
    heading.paragraph_format.space_after = Pt(18)
    _set_runs_font(heading, size_pt=14, bold=True, color=RGBColor(0x1C, 0x3D, 0x59))
    _format_body_paragraph(document.add_paragraph(
        "Este documento foi compilado automaticamente pelo SRA - Sistema de Relatórios de Atividades - "
        f"em {date.today().strftime('%d/%m/%Y')}, e está pronto para assinatura eletrônica via plataforma D4Sign."
    ))
    for text in ("Coordenação Técnica - Consórcio Concremat-Transplan", "Fiscalização - SEMIL/DER-SP"):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(35)
        paragraph.paragraph_format.space_after = Pt(30)
        _set_paragraph_top_border(paragraph, val="single", color="1C3D59", size="8")
        paragraph.add_run(text)
        _set_runs_font(paragraph, size_pt=9.5)
    note = document.add_paragraph(
        "Documento gerado a partir das submissões individuais da equipe multidisciplinar, consolidado pelo editor responsável. "
        f"Versão {rel.versao} · Código de referência: {rel.codigo} · Status: {rel.status}."
    )
    note.paragraph_format.space_before = Pt(30)
    note.paragraph_format.line_spacing = 1.6
    _set_runs_font(note, size_pt=8.5, color=RGBColor(0x4F, 0x5D, 0x6E))


def _add_texto(document: Document, texto: str) -> None:
    if not (texto or "").strip():
        return
    linhas = texto.splitlines()
    i = 0
    while i < len(linhas):
        raw = linhas[i]
        if not (raw or "").strip():
            i += 1
            continue
        line = raw.strip()
        if line.startswith("## "):
            _format_heading(document.add_heading(line[3:].strip(), level=3), 3)
        elif line.startswith("# "):
            _format_heading(document.add_heading(line[2:].strip(), level=2), 2)
        elif line_is_list_item(linhas[i]):
            j = i
            while j < len(linhas) and line_is_list_item(linhas[j]):
                j += 1
            _add_lista_bloco_docx(document, "\n".join(linhas[i:j]))
            i = j
            continue
        else:
            paragraph = document.add_paragraph(line)
            _format_body_paragraph(paragraph)
        i += 1


def _write_table_grid(document: Document, cells: list[list[str]]) -> None:
    cols = max(len(row) for row in cells)
    table = document.add_table(rows=len(cells), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(cells):
        for col_idx in range(cols):
            cell = table.cell(row_idx, col_idx)
            cell.text = row[col_idx] if col_idx < len(row) else ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Pt(0)
                paragraph_format.space_after = Pt(0)
                paragraph_format.line_spacing = 1.15
                _set_runs_font(paragraph, size_pt=8.8 if row_idx == 0 else 9.5, bold=row_idx == 0)


def _add_table(
    document: Document,
    conteudo: str,
    numero: str,
    opts: _DocxTabelaOpts | None = None,
) -> None:
    caption = opts or _DocxTabelaOpts()
    legenda = caption.legenda
    fonte = caption.fonte
    posicao = caption.posicao
    if legenda and posicao != "I":
        _format_caption(document.add_paragraph(f"Tabela {numero}: {legenda}"))
    linhas = [ln for ln in (conteudo or "").splitlines() if ln.strip()]
    linhas = [ln for ln in linhas if not re.fullmatch(r"-+(\s*\|\s*-+)*", ln.strip())]
    if linhas:
        cells = [[cell.strip() for cell in ln.strip().strip("|").split("|")] for ln in linhas]
        _write_table_grid(document, cells)
    if legenda and posicao == "I":
        _format_caption(document.add_paragraph(f"Tabela {numero}: {legenda}"))
    if fonte:
        _format_caption_fonte(document.add_paragraph(f"Fonte: {fonte}"))


def _add_figura(
    document: Document,
    figura: Figura | None,
    numero: str,
    opts: _DocxFiguraOpts | None = None,
) -> None:
    caption = opts or _DocxFiguraOpts()
    legenda = caption.legenda
    fonte = caption.fonte
    posicao = caption.posicao
    if legenda and posicao != "I":
        paragraph = document.add_paragraph(f"Figura {numero}: {legenda}")
        _format_caption(paragraph)
    if figura is not None:
        try:
            document.add_picture(BytesIO(figura.dados), width=Cm(17))
            image_paragraph = document.paragraphs[-1]
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_paragraph.paragraph_format.space_before = Pt(0)
            image_paragraph.paragraph_format.space_after = Pt(0)
        except Exception:  # noqa: BLE001
            _format_caption(document.add_paragraph(f"[Figura #{figura.id} não pôde ser inserida no DOCX]"))
    else:
        _format_caption(document.add_paragraph("[Figura importada sem imagem vinculada]"))
    if legenda and posicao == "I":
        paragraph = document.add_paragraph(f"Figura {numero}: {legenda}")
        _format_caption(paragraph)
    if fonte:
        paragraph = document.add_paragraph(f"Fonte: {fonte}")
        _format_caption_fonte(paragraph)


def _parse_figura_marker(match: re.Match[str]) -> tuple[str, int, str, str]:
    g1 = (match.group(1) or "").strip()
    g2 = match.group(2)
    g3 = match.group(3)
    g4 = match.group(4)
    idx_raw = ""
    figura_id = 0
    legenda = ""
    posicao = "I"
    if g4 is not None and g3 in ("S", "I"):
        idx_raw = g1
        figura_id = int(g2 or "0") if (g2 or "").isdigit() else 0
        posicao = g3
        legenda = (g4 or "").strip()
    elif g3 is not None and (g2 or "").isdigit():
        idx_raw = g1
        figura_id = int(g2 or "0")
        legenda = (g3 or "").strip()
    elif g2 is not None:
        figura_id = int(g1 or "0") if g1.isdigit() else 0
        legenda = (g2 or "").strip()
    else:
        figura_id = int(g1 or "0") if g1.isdigit() else 0
    return idx_raw, figura_id, legenda, posicao


def _parts_texto_e_tabelas(
    conteudo: str,
    sec_numero: str,
    tab_counter: int,
) -> tuple[list[tuple[str, str | tuple[str, str, str, str, str]]], int]:
    parts: list[tuple[str, str | tuple[str, str, str, str, str]]] = []
    last = 0
    text = conteudo or ""
    for match in _RE_TABELA.finditer(text):
        parts.append(("texto", text[last:match.start()]))
        idx_raw, posicao, legenda, corpo_tab = _parse_tabela_marker(match)
        tab_counter += 1
        numero = idx_raw or _figura_label(sec_numero, tab_counter)
        parts.append(
            (
                "tabela",
                (corpo_tab, legenda, "", numero, posicao),
            )
        )
        last = match.end()
    parts.append(("texto", text[last:]))
    return parts, tab_counter


def _emit_figuras_em_trecho_texto(
    document: Document,
    chunk: str,
    figuras_by_id: dict[int, Figura],
    sec_numero: str,
    contadores: _MarcadoresContadores,
) -> None:
    sub_last = 0
    for match in _RE_FIGURA.finditer(chunk):
        _add_texto(document, chunk[sub_last:match.start()])
        idx_raw, figura_id, legenda, posicao = _parse_figura_marker(match)
        contadores.fig += 1
        numero = idx_raw or _figura_label(sec_numero, contadores.fig)
        _add_figura(
            document,
            figuras_by_id.get(figura_id),
            numero,
            _DocxFiguraOpts(legenda=legenda, fonte="", posicao=posicao),
        )
        sub_last = match.end()
    _add_texto(document, chunk[sub_last:])


def _add_texto_com_marcadores(
    document: Document,
    conteudo: str,
    figuras_by_id: dict[int, Figura],
    sec_numero: str,
    contadores: _MarcadoresContadores,
) -> None:
    parts, contadores.tab = _parts_texto_e_tabelas(conteudo, sec_numero, contadores.tab)
    for kind, value in parts:
        if kind == "tabela":
            corpo, legenda, fonte, numero, posicao = value  # type: ignore[misc]
            _add_table(
                document,
                corpo,
                numero,
                _DocxTabelaOpts(legenda=legenda, fonte=fonte, posicao=posicao),
            )
            continue
        _emit_figuras_em_trecho_texto(document, str(value), figuras_by_id, sec_numero, contadores)


def _render_blocos_da_secao(
    document: Document,
    sec: Secao,
    figuras_by_id: dict[int, Figura],
    fig_counter: int,
    tab_counter: int,
) -> tuple[int, int]:
    for bloco in sec.blocos:
        if bloco.titulo:
            # PDF renderiza ``bloco.titulo`` como h2 12pt bold (regra
            # ``.bloco h2`` em ``app/templates/pdf/relatorio.html``).
            # Antes mapeávamos para H4 (10pt italic), o que apertava
            # visualmente. H2 = level 2 alinha o DOCX ao PDF.
            _format_heading(document.add_heading(bloco.titulo, level=2), 2)
        if bloco.tipo == "figura":
            fig_counter += 1
            _add_figura(
                document,
                figuras_by_id.get(bloco.figura_id or 0),
                _figura_label(sec.numero, fig_counter),
                _DocxFiguraOpts(legenda=bloco.legenda, fonte=bloco.fonte),
            )
        elif bloco.tipo == "tabela":
            tab_counter += 1
            _add_table(
                document,
                bloco.conteudo or "",
                _figura_label(sec.numero, tab_counter),
                _DocxTabelaOpts(legenda=bloco.legenda, fonte=bloco.fonte),
            )
        elif bloco.tipo == "lista":
            _add_lista_bloco_docx(document, bloco.conteudo or "")
        else:
            mc = _MarcadoresContadores(fig=fig_counter, tab=tab_counter)
            _add_texto_com_marcadores(
                document,
                bloco.conteudo or "",
                figuras_by_id,
                sec.numero,
                mc,
            )
            fig_counter = mc.fig
            tab_counter = mc.tab
    return fig_counter, tab_counter


def _render_secoes_corpo_docx(
    document: Document,
    rel: Relatorio,
    section_ids: set[int] | None,
    figuras_by_id: dict[int, Figura],
) -> None:
    fig_by_top: dict[str, int] = {}
    tab_by_top: dict[str, int] = {}
    body_started = False
    for sec in rel.secoes:
        if section_ids is not None and sec.id not in section_ids:
            continue
        top = (sec.numero or "").split(".")[0]
        fig_counter = fig_by_top.get(top, 0)
        tab_counter = tab_by_top.get(top, 0)
        heading_level = min(sec.numero.count(".") + 1, 6)
        if body_started and heading_level == 1:
            document.add_page_break()
        body_started = True
        _add_numbered_heading(document, sec.numero, sec.titulo, heading_level)
        if not sec.blocos:
            _format_empty_section(document.add_paragraph("— sem conteúdo nesta seção —"))
            continue
        fig_counter, tab_counter = _render_blocos_da_secao(
            document, sec, figuras_by_id, fig_counter, tab_counter,
        )
        fig_by_top[top] = fig_counter
        tab_by_top[top] = tab_counter


def render_docx(db: Session, rel: Relatorio, section_ids: set[int] | None = None) -> bytes:
    document = Document()
    _configure_document_styles(document)
    document.core_properties.title = rel.titulo or rel.codigo or "Relatório Mensal"
    document.core_properties.subject = f"{rel.codigo} - {rel.versao}"

    _add_cover(document, rel)
    _add_ficha_tecnica(document, rel)
    _add_sumario(document, rel, section_ids)

    header_paragraph = document.add_paragraph(rel.titulo or "Relatório Mensal")
    _format_heading(header_paragraph, 1)
    _format_body_paragraph(document.add_paragraph(f"{rel.codigo} - {rel.mes_referencia} - {rel.versao}"))
    _format_body_paragraph(document.add_paragraph(f"Período: {rel.periodo_inicio.strftime('%d/%m/%Y')} a {rel.periodo_fim.strftime('%d/%m/%Y')}"))

    figura_ids = _figura_ids_relatorio(rel, section_ids)
    figuras_by_id = {
        figura.id: figura for figura in db.query(Figura).filter(Figura.id.in_(figura_ids)).all()
    } if figura_ids else {}

    _render_secoes_corpo_docx(document, rel, section_ids, figuras_by_id)

    _add_assinaturas(document, rel)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
