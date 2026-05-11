# pylint: disable=protected-access,too-many-locals,too-many-branches,too-many-statements,unsubscriptable-object  # noqa: E501
"""Extracao de estrutura (secoes + blocos) a partir de relatorios DOCX.

Substitui a antiga extracao via PDF para a opcao "Relatorio entregue" no
formulario de criacao de relatorio (dashboard). DOCX preserva semantica
(heading, numeracao, listas, tabelas, imagens), o que fornece fidelidade
muito superior ao parsing de texto extraido de PDF.

Fluxo:
  1. ``listar_docx_disponiveis()`` lista DOCX em ``relatorios_entregues/``.
  2. ``extrair_relatorio_docx_disponivel(nome)`` abre o arquivo e retorna
     ``list[dict]`` com ``{secao_numero, secao_titulo, blocos:[...]}``
     compativel com o contrato ja consumido em ``app/routes/relatorios.py``
     (cria secoes + blocos em ``tx_session``).

Reaproveita utilitarios de ``app/routes/importacao.py`` (regex, iteracao
de blocos docx, listas numeradas Word) para manter uma fonte unica de
regras. Aqui nao ha ``Secao`` previa no banco: a arvore e inferida do
proprio DOCX.
"""

from __future__ import annotations

import base64
import re
import unicodedata
from html import escape as escape_html
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .list_lines import block_is_homogeneous_list, line_is_list_item
from .services.importacao import (
    _FIGURA_PREFIX_RE,
    _FIGURA_RE,
    _FONTE_RE,
    _HEADING_RE,
    _SECTION_NUMBER_RE,
    _TABELA_PREFIX_RE,
    _TABELA_RE,
    _get_w_numpr,
    _iter_docx_blocks,
    _paragraph_images,
    _read_numfmt_from_docx,
    _word_list_canonical_line,
)


def _norm_text_(s: str) -> str:
    """Normaliza texto para comparação: strip, lower, sem acentos, espaços colapsados."""  # noqa: E501
    if not s:
        return ""
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    s = s.lower()
    s = re.sub(r"\s+", " ", s).strip()
    return s


PASTA_RELATORIOS = Path(__file__).resolve().parent.parent / "relatorios_entregues"  # noqa: E501


def listar_docx_disponiveis() -> list[str]:
    """Lista nomes de DOCX em ``relatorios_entregues/`` (ordenados).

    Ignora arquivos temporarios do Word (``~$...``) e extensoes que nao
    sejam ``.docx``.
    """
    if not PASTA_RELATORIOS.is_dir():
        return []
    return sorted(
        p.name
        for p in PASTA_RELATORIOS.iterdir()
        if p.suffix.lower() == ".docx" and not p.name.startswith("~$")  # noqa: E501
    )


_STYLE_HEADING_LEVEL_RE = re.compile(r"(\d+)")


def _style_heading_level(style_name: str) -> int:
    """Devolve o nivel de heading (1..9) a partir do nome do estilo Word.

    Funciona com ``Heading 1``, ``heading 2``, ``Titulo 1``, ``Título 3``,
    ``Cabecalho 4``, etc. Devolve 0 quando o paragrafo nao e heading.
    """
    s = (style_name or "").strip().lower()
    if not s:
        return 0
    if not (
        s.startswith("heading")
        or s.startswith("título")
        or s.startswith("titulo")
        or s.startswith("cabeçalho")
        or s.startswith("cabecalho")
    ):
        return 0
    m = _STYLE_HEADING_LEVEL_RE.search(s)
    if not m:
        return 1
    try:
        lvl = int(m.group(1))
    except ValueError:
        return 1
    return max(1, min(lvl, 9))


def _split_legenda_fonte(line: str) -> tuple[str, str]:
    match = _FONTE_RE.search(line)
    if not match:
        return line.strip(), ""
    return line[: match.start()].strip(), match.group(1).strip()


def _split_figura_fonte(line: str) -> tuple[str, str]:
    legenda, fonte = _split_legenda_fonte(line)
    legenda = _FIGURA_PREFIX_RE.sub("", legenda).strip()
    return legenda, fonte


def _split_tabela_fonte(line: str) -> tuple[str, str]:
    legenda, fonte = _split_legenda_fonte(line)
    legenda = _TABELA_PREFIX_RE.sub("", legenda).strip()
    return legenda, fonte


def _parse_numero_titulo(text: str) -> tuple[str, str]:
    """Extrai ``(numero, titulo)`` de uma linha de heading.

    Exemplos aceitos:
      - "4.4.1 Acompanhamento tecnico"  -> ("4.4.1", "Acompanhamento tecnico")
      - "4 Visao geral"                 -> ("4", "Visao geral")
      - "4.1 - Coordenacao"             -> ("4.1", "Coordenacao")
      - "Apresentacao"                  -> ("", "Apresentacao")
    """
    body = re.sub(r"^#{1,6}\s*", "", (text or "").strip())
    if not body:
        return "", ""
    m = _SECTION_NUMBER_RE.match(body)
    if not m:
        return "", body
    numero = (m.group(1) or "").strip()
    resto = (m.group(2) or "").strip()
    resto = re.sub(r"^[\s\-\u2013\u2014.:]+", "", resto).strip()
    return numero, resto


def _proximo_numero_por_nivel(contadores: dict, nivel: int, ultimo_numero: str) -> str:  # noqa: E501
    """Gera um numero de secao quando o heading nao tras numero explicito.

    Baseado nos contadores por nivel (``contadores[nivel]``) e no ultimo
    numero emitido, para formar uma cadeia ``X.Y.Z`` coerente.
    """
    nivel = max(1, nivel)
    # Extrai prefixo do ultimo numero ate nivel-1
    partes_prefixo: list[str] = []
    if nivel > 1 and ultimo_numero:
        partes = [p for p in ultimo_numero.split(".") if p]
        partes_prefixo = partes[: nivel - 1]
        # Se o ultimo numero e mais curto que nivel-1, completa com 1
        while len(partes_prefixo) < nivel - 1:
            partes_prefixo.append("1")
    # Limpa contadores de niveis mais profundos
    for k in list(contadores.keys()):
        if k > nivel:
            del contadores[k]
    contadores[nivel] = contadores.get(nivel, 0) + 1
    partes_prefixo.append(str(contadores[nivel]))
    return ".".join(partes_prefixo)


def _abrir_secao(
    resultado: list,
    numero: str,
    titulo: str,
    vistos: set,
    orientacao: str = "portrait",
) -> dict | None:
    """Cria uma nova entrada de secao e devolve-a, ou ``None`` se duplicada/invalida."""  # noqa: E501
    numero = (numero or "").strip()
    titulo = (titulo or "").strip()
    if not numero or numero in vistos or len(numero) > 16:
        return None
    vistos.add(numero)
    sec = {
        "secao_numero": numero,
        "secao_titulo": titulo or f"Secao {numero}",
        "orientacao": orientacao if orientacao in ("portrait", "landscape") else "portrait",
        "blocos": [],
    }
    resultado.append(sec)
    return sec


def _append_tabela_para_secao(sec: dict, linhas: list, legenda: str, fonte: str) -> None:  # noqa: E501
    if sec is None or not any((ln or "").strip() for ln in linhas):
        return
    conteudo = "\n".join(ln for ln in linhas if (ln or "").strip())
    sec["blocos"].append(
        {
            "tipo": "tabela",
            "conteudo": conteudo,
            "legenda": (legenda or "").strip(),
            "fonte": (fonte or "").strip(),
        }
    )


def _b64_to_bytes(b64_str: str) -> bytes:
    """Converte string base64 para bytes."""
    if not b64_str:
        return b""
    return base64.b64decode(b64_str, validate=True)


def _append_figura_para_secao(sec: dict, legenda: str, fonte: str, dados: bytes | None, mime: str) -> int | None:  # noqa: E501
    """Adiciona figura à seção e retorna o índice do bloco."""
    if sec is None:
        return None
    idx = len(sec["blocos"])
    sec["blocos"].append(
        {
            "tipo": "figura",
            "legenda": legenda.strip(),
            "fonte": fonte.strip(),
            "dados": dados,
            "dados_imagem": dados,
            "mime": mime,
        }
    )
    return idx


def _eh_ruido_modelo(text: str) -> bool:
    """Verifica se texto é ruído de modelos legados."""
    text = text.strip()
    padroes_ruido = [
        r"^bloco:\s*paragrafo\s*\(texto\)$",
        r"^figura\s+\d+\.?\d*:\s*$",
        r"^tabela\s+\d+\.?\d*:\s*$",
        r"^---+$",
        r"^\.{3,}$",
    ]
    for padrao in padroes_ruido:
        if re.match(padrao, text, re.IGNORECASE):
            return True
    return False


# ---------------------------------------------------------------------------
# Extração de formatação (runs → HTML)
# ---------------------------------------------------------------------------

RICH_HTML_MARKER = "<!--SRA_RICH-->"
_CALLOUT_APRESENTACAO_RE = re.compile(
    r"este relat[oó]rio\s+[ée]\s+um item da medi[cç][aã]o\s+\d+",
    re.IGNORECASE,
)


def _orientacao_da_secao_docx(document: Document, section_idx: int) -> str:
    """Devolve 'landscape' ou 'portrait' lendo ``w:pgSz/w:orient`` do DOCX.

    No OOXML, o documento é dividido em "sections" (``w:sectPr``).
    ``document.sections[i].orientation`` expõe ``WD_ORIENT.PORTRAIT`` /
    ``LANDSCAPE``. Esta função é segura mesmo quando o índice está fora
    do alcance ou o atributo não está presente.
    """
    try:
        sections = document.sections
        if section_idx >= len(sections):
            return "portrait"
        sec = sections[section_idx]
        # python-docx: 0=portrait, 1=landscape
        return "landscape" if int(sec.orientation) == 1 else "portrait"
    except (AttributeError, ValueError, IndexError):
        return "portrait"


def _paragraph_section_index(paragraph: Paragraph, doc_xml: list) -> int:
    """Calcula o índice da `sectPr` cujo conteúdo este parágrafo integra.

    Itera linearmente os elementos do body; cada `sectPr` filho de `pPr`
    (final de section) ou direto de `body` fecha a section corrente.
    Retorna 0 antes do primeiro fechamento.
    """
    p_el = paragraph._element
    idx = 0
    for el in doc_xml:
        if el is p_el:
            return idx
        if el.tag.endswith("}p"):
            # uma sectPr dentro do pPr deste parágrafo encerra a section
            ppr = el.find(qn("w:pPr"))
            if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
                idx += 1
        elif el.tag.endswith("}sectPr"):
            idx += 1
    return idx


def _run_to_html(run) -> str:
    """Converte um ``Run`` python-docx em HTML inline (negrito, itálico, cor, etc.)."""  # noqa: E501
    from docx.text.run import Run

    if not isinstance(run, Run):
        return escape_html(run.text) if hasattr(run, "text") else str(run)

    text = run.text or ""
    if not text:
        return ""

    # Evita quebrar espaços significativos no HTML
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    tags: list[str] = []
    styles: list[str] = []

    if run.bold:
        tags.append("strong")
    if run.italic:
        tags.append("em")
    if run.underline:
        tags.append("u")
    if getattr(run.font, "strike", False):
        tags.append("s")
    if run.font.subscript:
        tags.append("sub")
    if run.font.superscript:
        tags.append("sup")

    color = run.font.color
    if color and color.rgb:
        styles.append(f"color:#{color.rgb}")
    if run.font.size:
        pt = run.font.size.pt
        styles.append(f"font-size:{pt:.1f}pt")
    if run.font.name:
        styles.append(f"font-family:{run.font.name}")
    if run.font.highlight_color:
        # python-docx não expõe cor exata de highlight; pulamos
        pass

    out = text
    if styles:
        style_attr = ";".join(styles)
        out = f'<span style="{style_attr}">{out}</span>'
    for tag in reversed(tags):
        out = f"<{tag}>{out}</{tag}>"
    return out


def _paragraph_alignment_style(paragraph) -> str:
    """Devolve CSS ``text-align`` a partir do alinhamento do parágrafo."""
    al = paragraph.paragraph_format.alignment
    if al == WD_ALIGN_PARAGRAPH.CENTER:
        return "center"
    if al == WD_ALIGN_PARAGRAPH.RIGHT:
        return "right"
    if al == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return "justify"
    if al == WD_ALIGN_PARAGRAPH.LEFT:
        return "left"
    return ""


def _paragraph_format_styles(paragraph) -> list[str]:
    """CSS inline para recuo/espaçamento do parágrafo (lista de regras)."""
    out: list[str] = []
    pf = paragraph.paragraph_format
    if pf.left_indent is not None:
        out.append(f"margin-left:{pf.left_indent.pt:.1f}pt")
    if pf.first_line_indent is not None:
        out.append(f"text-indent:{pf.first_line_indent.pt:.1f}pt")
    if pf.space_after is not None:
        out.append(f"margin-bottom:{pf.space_after.pt:.1f}pt")
    if pf.space_before is not None:
        out.append(f"margin-top:{pf.space_before.pt:.1f}pt")
    if pf.line_spacing is not None and isinstance(pf.line_spacing, (int, float)):  # noqa: E501
        out.append(f"line-height:{pf.line_spacing:.2f}")
    return out


def _paragraph_to_html(paragraph) -> str:
    """Converte um Paragraph python-docx em fragmento HTML de paragrafo."""
    inner = "".join(_run_to_html(r) for r in paragraph.runs)
    if not inner:
        inner = escape_html(paragraph.text or "")

    styles: list[str] = []
    ta = _paragraph_alignment_style(paragraph)
    if ta:
        styles.append(f"text-align:{ta}")
    styles.extend(_paragraph_format_styles(paragraph))

    if styles:
        p_html = f'<p style="{";".join(styles)}">{inner}</p>'  # noqa: H025
    else:
        p_html = f"<p>{inner}</p>"
    if _CALLOUT_APRESENTACAO_RE.search(paragraph.text or ""):
        return f'<div class="sra-docx-callout">{p_html}</div>'
    return p_html


# ---------------------------------------------------------------------------
# Extração de tabelas (→ HTML rico)
# ---------------------------------------------------------------------------


def _tc_span_info(tc) -> tuple[int, int]:
    """Devolve (colspan, rowspan) a partir do XML de uma célula."""
    grid_span = 1
    vmerge_start = False
    vmerge_continue = False

    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is not None:
        grid_span_el = tc_pr.find(qn("w:gridSpan"))
        if grid_span_el is not None:
            val = grid_span_el.get(qn("w:val"))
            if val:
                try:
                    grid_span = int(val)
                except ValueError:
                    pass
        vmerge_el = tc_pr.find(qn("w:vMerge"))
        if vmerge_el is not None:
            vm_val = vmerge_el.get(qn("w:val"))
            if vm_val == "restart":
                vmerge_start = True
            else:
                vmerge_continue = True

    if vmerge_start:
        # Conta quantas células abaixo continuam o merge
        rowspan = 1
    elif vmerge_continue:
        return 0, 0  # célula mesclada verticalmente → omitir no HTML
    else:
        rowspan = 1

    return grid_span, rowspan


# pylint: disable-next=too-many-nested-blocks
def _cell_borders_and_bg(tc) -> list[str]:  # noqa: E501
    """Extrai estilos CSS de bordas e cor de fundo da célula."""
    styles: list[str] = []
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return styles

    shd = tc_pr.find(qn("w:shd"))
    if shd is not None:
        fill = shd.get(qn("w:fill"))
        if fill and fill != "auto" and fill != "000000":
            styles.append(f"background-color:#{fill}")

    tc_borders = tc_pr.find(qn("w:tcBorders"))
    # pylint: disable=too-many-nested-blocks
    if tc_borders is not None:
        for edge in ("top", "left", "bottom", "right"):
            edge_el = tc_borders.find(qn(f"w:{edge}"))
            if edge_el is not None:
                sz = edge_el.get(qn("w:sz"))
                color = edge_el.get(qn("w:color"))
                val = edge_el.get(qn("w:val"))
                if val and val != "nil" and val != "none":
                    css_edge = {
                        "top": "border-top",
                        "left": "border-left",
                        "bottom": "border-bottom",
                        "right": "border-right",
                    }[edge]
                    width = "1pt"
                    if sz:
                        try:
                            width = f"{int(sz) / 8:.1f}pt"
                        except ValueError:
                            pass
                    c = f"#{color}" if color and color != "auto" else "#000000"
                    styles.append(f"{css_edge}:{width} solid {c}")
    return styles


def _cell_to_html(cell) -> str:
    """Converte célula python-docx em elemento HTML `td` com formatação."""
    tc = cell._tc
    colspan, rowspan = _tc_span_info(tc)
    if colspan == 0 and rowspan == 0:
        return ""  # célula mesclada verticalmente

    attrs: list[str] = []
    if colspan > 1:
        attrs.append(f'colspan="{colspan}"')
    if rowspan > 1:
        attrs.append(f'rowspan="{rowspan}"')

    # Alinhamento vertical
    v_align = ""
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is not None:
        v_align_el = tc_pr.find(qn("w:vAlign"))
        if v_align_el is not None:
            v = v_align_el.get(qn("w:val"))
            if v:
                v_align = v

    styles: list[str] = []
    if v_align:
        styles.append(f"vertical-align:{v_align}")
    styles.extend(_cell_borders_and_bg(tc))

    # Alinhamento horizontal da célula via parágrafo(s)
    for para in cell.paragraphs:
        ta = _paragraph_alignment_style(para)
        if ta:
            styles.append(f"text-align:{ta}")
            break

    if styles:
        attrs.append(f'style="{";".join(styles)}"')

    attr_str = " ".join(attrs)
    tag_open = f"<td {attr_str}>" if attr_str else "<td>"

    # Conteúdo dos parágrafos da célula
    inner_parts: list[str] = []
    for para in cell.paragraphs:
        inner_parts.append(_paragraph_to_html(para))
    inner = "".join(inner_parts)
    return f"{tag_open}{inner}</td>"


def _table_to_html(table) -> str:
    """Converte uma ``Table`` python-docx em ``table`` HTML."""
    rows_html: list[str] = []
    for row in table.rows:
        cells_html = [_cell_to_html(cell) for cell in row.cells if _cell_to_html(cell)]  # noqa: E501
        if cells_html:
            rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
    if not rows_html:
        return ""
    return "<table>" + "".join(rows_html) + "</table>"


# ---------------------------------------------------------------------------
# Buffer de blocos (agora suporta HTML rico)
# ---------------------------------------------------------------------------

RICH_HTML_MARKER = "<!--SRA_RICH-->"


def _flush_buffer_texto(sec: dict, buf: list[str]) -> None:
    """Consome o buffer de texto e gera bloco ``texto`` ou ``lista``.

    Se detectar HTML no buffer (marcadores RICH ou tags), grava como bloco
    ``texto`` com conteúdo rico (prefixado por ``RICH_HTML_MARKER``).
    Quando o buffer for uma lista e contiver HTML, converte para
    ul/ol (lista nao ordenada/ordenada)."""
    if sec is None:
        buf.clear()
        return
    clean = [ln.rstrip() for ln in buf if ln.strip()]
    buf.clear()
    if not clean:
        return

    # Se qualquer linha tiver tag HTML ou já for rich, trata tudo como HTML
    has_html = any(  # noqa: E501
        (ln.startswith("<") and ">" in ln) or RICH_HTML_MARKER in ln  # noqa: E501
        for ln in clean
    )

    if has_html:
        is_list = all(  # noqa: E501
            line_is_list_item(ln) or (ln.startswith("<") and "</li>" in ln)  # noqa: E501,H025
            for ln in clean
        )
        if is_list:
            joined = RICH_HTML_MARKER + "\n" + _lista_para_html(clean)
        else:
            joined = "".join(clean)
            if not joined.startswith(RICH_HTML_MARKER):
                joined = RICH_HTML_MARKER + "\n" + joined
        sec["blocos"].append(
            {
                "tipo": "texto",
                "conteudo": joined,
            }
        )
        return

    tipo = "lista" if block_is_homogeneous_list(clean) else "texto"
    sec["blocos"].append(
        {
            "tipo": tipo,
            "conteudo": "\n".join(clean),
        }
    )


def _paragraph_to_simple_or_html(paragraph) -> str:
    """Devolve texto simples se não houver formatação especial;
    caso contrário, HTML."""
    html = _paragraph_to_html(paragraph)
    text = (paragraph.text or "").strip()
    if not text:
        return ""
    # Remove <p ...> e </p> para verificar se há tags internas  # noqa: H025
    cleaned = re.sub(r"^<p[^>]*>", "", html)
    cleaned = re.sub(r"</p>$", "", cleaned)
    if "<" not in cleaned:
        return text
    return html


def _lista_para_html(buf: list[str]) -> str:
    """Converte buffer de itens de lista (alguns podem ser itens li HTML)
    em lista nao ordenada ou ordenada."""
    if not buf:
        return ""
    # Detecta se há itens numerados
    has_ol = any(
        re.match(r"^\d+[.)]\s+", item) or item.startswith("<li>") and re.search(r"^\d+[.)]", item)  # noqa: E501
        for item in buf  # noqa: E501
    )
    tag = "ol" if has_ol else "ul"
    items: list[str] = []
    for item in buf:
        if item.startswith("<") and "</li>" in item:
            items.append(item)
        else:
            body = item
            # Remove prefixos de lista
            m = re.match(r"^(?:[-*•]|\d+[.)]|[a-zA-Z][.)]|\([\diI]\))\s+(.*)$", body)  # noqa: E501
            if m:
                body = m.group(1)
            items.append(f"<li>{body}</li>")
    return f"<{tag}>{''.join(items)}</{tag}>"


def extrair_relatorio_docx(raw: bytes) -> list[dict]:  # noqa: C901
    """Extrai estrutura completa (secoes + blocos) de bytes de um DOCX.

    Retorno: ``list[dict]`` onde cada item tem as chaves
    ``secao_numero``, ``secao_titulo`` e ``blocos``. Cada bloco tem
    ``tipo`` (``texto`` | ``lista`` | ``tabela`` | ``figura``) e campos
    dependentes do tipo. Formato alinhado ao consumido pela rota de
    criacao de relatorio (``app/routes/relatorios.py``).
    """
    document = Document(BytesIO(raw))

    # Cache da lista de elementos do body para resolver índice de section
    # a que cada parágrafo pertence (espelha w:sectPr → orientação).
    _body_children = list(document.element.body.iterchildren())

    resultado: list[dict] = []
    vistos_numeros: set[str] = set()
    contadores_por_nivel: dict[int, int] = {}
    ultimo_numero = ""
    current_sec: dict | None = None
    buf: list[str] = []
    word_list_counters: dict[tuple[int, int], int] = {}

    last_media_kind = ""  # "figura" ou "tabela"
    last_media_idx: int | None = None
    pending_figure_idx: int | None = None
    pending_table_legenda = ""
    pending_table_fonte = ""

    def _sec_ref() -> dict | None:
        return current_sec

    for element in _iter_docx_blocks(document):
        # ------------------- Tabelas nativas -------------------
        if isinstance(element, Table):
            _flush_buffer_texto(_sec_ref(), buf)
            tabela_html = _table_to_html(element)
            if tabela_html and current_sec is not None:
                rich = RICH_HTML_MARKER + "\n" + tabela_html
                current_sec["blocos"].append(
                    {
                        "tipo": "texto",
                        "conteudo": rich,
                        "legenda": (pending_table_legenda or "").strip(),
                        "fonte": (pending_table_fonte or "").strip(),
                    }
                )
                last_media_idx = len(current_sec["blocos"]) - 1
                last_media_kind = "tabela"
            pending_table_legenda = ""
            pending_table_fonte = ""
            continue

        # ------------------- Paragrafos -------------------
        paragraph: Paragraph = element

        # Imagens incorporadas no paragrafo (placeholder; legenda vira depois)
        imgs = _paragraph_images(paragraph)
        if imgs:
            _flush_buffer_texto(_sec_ref(), buf)
            pending_figure_idx = None
            for img in imgs:
                dados = _b64_to_bytes(img.get("image_b64") or "")
                mime = img.get("image_mime") or "image/png"
                pending_figure_idx = _append_figura_para_secao(_sec_ref(), "", "", dados, mime)
                if pending_figure_idx is not None:
                    last_media_idx = pending_figure_idx
                    last_media_kind = "figura"

        text = (paragraph.text or "").strip()

        # Filtra parágrafos cujo conteúdo é ruído de modelos legados
        # (ex.: "Bloco: paragrafo (texto)", "Figura X.Y: ...", "---", etc.).
        if text and _eh_ruido_modelo(text):
            continue

        # Fonte solta apos figura/tabela
        if text.lower().startswith("fonte:") and current_sec is not None and last_media_idx is not None:  # noqa: E501
            current_sec["blocos"][last_media_idx]["fonte"] = text[6:].strip()
            continue

        # Linha "Figura X.Y: ..." - legenda para figura pendente ou nova
        if text and _FIGURA_RE.match(text):
            _flush_buffer_texto(_sec_ref(), buf)
            legenda, fonte = _split_figura_fonte(text)
            if current_sec is not None:
                if pending_figure_idx is not None:
                    current_sec["blocos"][pending_figure_idx]["legenda"] = legenda  # noqa: E501
                    current_sec["blocos"][pending_figure_idx]["fonte"] = fonte
                    last_media_idx = pending_figure_idx
                    last_media_kind = "figura"
                    pending_figure_idx = None
                else:
                    # Figura citada sem binario - registra placeholder vazio
                    idx = _append_figura_para_secao(_sec_ref(), legenda, fonte, None, "image/png")  # noqa: E501
                    if idx is not None:
                        last_media_idx = idx
                        last_media_kind = "figura"
            continue

        # Linha "Tabela X.Y: ..." - legenda para tabela que vem em seguida
        if text and _TABELA_RE.match(text):
            _flush_buffer_texto(_sec_ref(), buf)
            legenda, fonte = _split_tabela_fonte(text)
            if (
                current_sec is not None
                and last_media_kind == "tabela"
                and last_media_idx is not None
                and not current_sec["blocos"][last_media_idx].get("legenda")
            ):
                current_sec["blocos"][last_media_idx]["legenda"] = legenda
                current_sec["blocos"][last_media_idx]["fonte"] = fonte
            else:
                pending_table_legenda = legenda
                pending_table_fonte = fonte
            continue

        # Heading (por estilo Word)
        style_name = ""
        try:
            style_name = (paragraph.style.name or "") if paragraph.style else ""  # noqa: E501
        except AttributeError:
            style_name = ""
        heading_level = _style_heading_level(style_name)

        # Heading detectado por regex textual (fallback quando estilo nao
        # e Heading/Titulo mas a linha tem "N.N titulo").
        if heading_level == 0 and text and _HEADING_RE.match(text):
            # Usa o numero como nivel aproximado (conta segmentos).
            numero_preview, _ = _parse_numero_titulo(text)
            if numero_preview:
                heading_level = len([p for p in numero_preview.split(".") if p])  # noqa: E501

        if heading_level > 0 and text:
            _flush_buffer_texto(_sec_ref(), buf)
            numero, titulo = _parse_numero_titulo(text)
            if not numero:
                numero = _proximo_numero_por_nivel(contadores_por_nivel, heading_level, ultimo_numero)  # noqa: E501
            else:
                # Alinha contadores ao numero explicito
                partes = [int(p) for p in numero.split(".") if p.isdigit()]
                for i, v in enumerate(partes, start=1):
                    contadores_por_nivel[i] = v
                for k in list(contadores_por_nivel.keys()):
                    if k > len(partes):
                        del contadores_por_nivel[k]
            # Detecta orientação da página em que este heading aparece.
            # Espelha w:pgSz/w:orient do DOCX (referência D20-13 usa
            # landscape em cronogramas / sec. 7, 11).
            sect_idx = _paragraph_section_index(paragraph, _body_children)
            orientacao = _orientacao_da_secao_docx(document, sect_idx)
            nova = _abrir_secao(resultado, numero, titulo, vistos_numeros, orientacao)
            if nova is not None:
                current_sec = nova
                ultimo_numero = numero
                pending_figure_idx = None
                last_media_idx = None
                last_media_kind = ""
            continue

        if not text:
            _flush_buffer_texto(_sec_ref(), buf)
            continue

        p_html_or_text = _paragraph_to_simple_or_html(paragraph)
        is_rich = p_html_or_text != text

        # Lista numerada/bullet nativa Word
        w_num = _get_w_numpr(paragraph)
        if w_num is not None:
            il, nid = w_num
            numfmt_word = _read_numfmt_from_docx(document, nid, il)
            if buf and not all(line_is_list_item(ln) for ln in buf):
                _flush_buffer_texto(_sec_ref(), buf)
            if is_rich:
                buf.append(f"<li>{p_html_or_text}</li>")
            else:
                buf.append(_word_list_canonical_line(text, il, numfmt_word, nid, word_list_counters))  # noqa: E501
            continue

        # Estilo de lista baseado em nome ("List Paragraph", "Lista")
        style_lower = (style_name or "").lower()
        if "list" in style_lower or "lista" in style_lower:
            if buf and not all(line_is_list_item(ln) for ln in buf):
                _flush_buffer_texto(_sec_ref(), buf)
            if is_rich:
                buf.append(f"<li>{p_html_or_text}</li>")
            else:
                buf.append("- " + text.lstrip("-\u2022\u00b7 "))
            continue

        # Texto corrido
        if buf and all(line_is_list_item(ln) for ln in buf):
            _flush_buffer_texto(_sec_ref(), buf)
        buf.append(p_html_or_text if is_rich else text)

    _flush_buffer_texto(_sec_ref(), buf)

    # Ordena pela chave do numero (estavel)
    def _sort_key(sec: dict) -> list:
        return [int(x) if x.isdigit() else 0 for x in (sec.get("secao_numero") or "").split(".")]  # noqa: E501

    resultado.sort(key=_sort_key)
    return resultado


def extrair_relatorio_docx_disponivel(nome_arquivo: str) -> list[dict]:
    """Extrai secoes + blocos de um DOCX da pasta ``relatorios_entregues/``.

    Valida ``nome_arquivo`` contra ``listar_docx_disponiveis()`` para
    evitar path traversal.
    """
    disponiveis = set(listar_docx_disponiveis())
    if nome_arquivo not in disponiveis:
        raise ValueError(f"DOCX nao disponivel: {nome_arquivo}")
    caminho = PASTA_RELATORIOS / nome_arquivo
    with open(caminho, "rb") as fh:
        raw = fh.read()
    return extrair_relatorio_docx(raw)
