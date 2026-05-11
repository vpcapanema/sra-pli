"""HTML rico (blocos marcados ``<!--SRA_RICH-->``) → trechos python-docx."""

from __future__ import annotations

import importlib
import re
from types import ModuleType

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from . import sra_rich_html as srh


def _lazy_docx_render() -> ModuleType:
    """Carrega ``docx_render`` em tempo de execução (evita ciclo estático)."""
    return importlib.import_module(".docx_render", __package__)


def _quill_para_alignment(classes) -> WD_ALIGN_PARAGRAPH | None:
    if not classes:
        return None
    seq = classes if isinstance(classes, list) else str(classes).split()
    for c in seq:
        if c == "ql-align-center":
            return WD_ALIGN_PARAGRAPH.CENTER
        if c == "ql-align-right":
            return WD_ALIGN_PARAGRAPH.RIGHT
        if c == "ql-align-justify":
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        if c == "ql-align-left":
            return WD_ALIGN_PARAGRAPH.LEFT
    return None


def _alignment_from_style(style_val: str | None) -> WD_ALIGN_PARAGRAPH | None:
    if not style_val:
        return None
    style = str(style_val).lower().replace(" ", "")
    mat = re.search(r"text-align:(left|center|right|justify)", style)
    if not mat:
        return None
    key = mat.group(1)
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(key)


def _paragraph_alignment_from_html_node(node) -> WD_ALIGN_PARAGRAPH | None:
    """Alinhamento a partir de classes Quill ou ``style`` do RTE."""
    al = _quill_para_alignment(node.get("class"))
    if al is not None:
        return al
    return _alignment_from_style(node.get("style"))


def _parse_span_style(style_val: str | None) -> dict[str, str]:
    """Extrai ``color``, ``font-size`` e ``font-family`` de ``style``."""
    out: dict[str, str] = {}
    if not style_val:
        return out
    style = str(style_val).lower().replace(" ", "")
    m = re.search(r"color:#?([0-9a-fA-F]{6})", style)
    if m:
        out["color"] = m.group(1)
    m = re.search(r"font-size:([\d.]+)pt", style)
    if m:
        out["size_pt"] = m.group(1)
    m = re.search(r"font-family:([^;]+)", style)
    if m:
        out["font_family"] = m.group(1).strip()
    return out


# pylint: disable-next=too-many-locals,too-many-branches
def _add_html_runs(
    paragraph,
    node,
    *,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    strike: bool = False,
    sub: bool = False,
    sup: bool = False,
    color: str | None = None,
    size_pt: float | None = None,
    font_name: str | None = None,
) -> None:
    """Adiciona runs ao parágrafo percorrendo a árvore HTML recursivamente."""
    from bs4 import NavigableString, Tag

    for child in node.children:
        if isinstance(child, NavigableString):
            t = str(child)
            if t:
                run = paragraph.add_run(t)
                run.bold = bold
                run.italic = italic
                run.underline = underline
                run.font.strike = strike
                run.font.subscript = sub
                run.font.superscript = sup
                if color:
                    try:
                        run.font.color.rgb = RGBColor.from_string(color)
                    except (ValueError, TypeError):
                        pass
                if size_pt is not None:
                    run.font.size = Pt(size_pt)
                if font_name:
                    run.font.name = font_name
        elif isinstance(child, Tag):
            name = child.name.lower() if child.name else ""
            new_bold = bold or name in ("strong", "b")
            new_italic = italic or name in ("em", "i")
            new_underline = underline or name == "u"
            new_strike = strike or name in ("s", "strike")
            new_sub = sub or name == "sub"
            new_sup = sup or name == "sup"
            new_color = color
            new_size = size_pt
            new_font = font_name
            if name == "span":
                parsed = _parse_span_style(child.get("style"))
                new_color = parsed.get("color", color)
                if "size_pt" in parsed:
                    try:
                        new_size = float(parsed["size_pt"])
                    except ValueError:
                        pass
                new_font = parsed.get("font_family", font_name)
            elif name == "a":
                href = child.get("href") or ""
                if href:
                    # python-docx não suporta hyperlink simples;
                    # incluímos o texto normal e, se possível, o URL
                    pass
            _add_html_runs(
                paragraph,
                child,
                bold=new_bold,
                italic=new_italic,
                underline=new_underline,
                strike=new_strike,
                sub=new_sub,
                sup=new_sup,
                color=new_color,
                size_pt=new_size,
                font_name=new_font,
            )


def _fill_paragraph_from_html_node(paragraph, node) -> None:
    """Preenche parágrafo a partir de nós filhos (ênfase, estilos inline)."""
    _add_html_runs(paragraph, node)


def _docx_rich_blockquote(document: Document, child) -> None:
    dr = _lazy_docx_render()
    quote_para = document.add_paragraph(child.get_text(strip=True))
    quote_para.paragraph_format.left_indent = Pt(18)
    quote_para.paragraph_format.space_before = Pt(6)
    dr.format_docx_body_paragraph(quote_para)


def _docx_rich_img_placeholder(document: Document, child) -> None:
    dr = _lazy_docx_render()
    src = child.get("src") or ""
    alt = (child.get("alt") or "").strip()
    legend = f"[Imagem: {alt}]" if alt else "[Imagem]"
    para = document.add_paragraph(f"{legend} ({src})")
    dr.format_docx_body_paragraph(para)


def _docx_rich_table(document: Document, child) -> None:
    rows_el = child.find_all("tr")
    if not rows_el:
        return
    col_counts = [len(tr.find_all(["td", "th"])) for tr in rows_el]
    max_cols = max(col_counts) if col_counts else 0
    if max_cols <= 0:
        return
    table = document.add_table(rows=len(rows_el), cols=max_cols)
    table.style = "Table Grid"
    for ri, tr in enumerate(rows_el):
        cells = tr.find_all(["td", "th"])
        for ci in range(max_cols):
            cell = table.rows[ri].cells[ci]
            if ci < len(cells):
                cell.text = cells[ci].get_text(strip=True)
            else:
                cell.text = ""


def _add_rich_html_node(document: Document, child) -> None:
    """Um nó de topo filho do wrapper ``div`` (HTML rico → DOCX)."""
    dr = _lazy_docx_render()

    if getattr(child, "name", None) is None:
        return
    name = child.name.lower()
    if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
        lvl = int(name[1])
        dr.format_docx_heading(
            document.add_heading(child.get_text(strip=True), level=lvl),
            lvl,
        )
    elif name == "blockquote":
        _docx_rich_blockquote(document, child)
    elif name == "img":
        _docx_rich_img_placeholder(document, child)
    elif name == "table":
        _docx_rich_table(document, child)
    elif name == "p":
        para = document.add_paragraph()
        _fill_paragraph_from_html_node(para, child)
        al = _paragraph_alignment_from_html_node(child)
        if al is not None:
            para.alignment = al
        dr.format_docx_body_paragraph(para)
    elif name in ("ul", "ol"):
        style = "List Bullet" if name == "ul" else "List Number"
        for li in child.find_all("li", recursive=False):
            item_para = document.add_paragraph(style=style)
            inner = li.find("p", recursive=False)
            src = inner if inner is not None else li
            _fill_paragraph_from_html_node(item_para, src)
            al = _paragraph_alignment_from_html_node(src)
            if al is not None:
                item_para.alignment = al
            dr.format_docx_body_paragraph(item_para)


def add_rich_html_docx(document: Document, html: str) -> None:
    """DOCX a partir de fragmento HTML sanitizado."""
    from bs4 import BeautifulSoup

    clean = srh.sanitize_rich_html_for_pdf(html)
    if not (clean or "").strip():
        return
    soup = BeautifulSoup("<div>" + clean + "</div>", "html.parser")
    root = soup.div
    if root is None:
        return
    for child in root.children:
        _add_rich_html_node(document, child)
