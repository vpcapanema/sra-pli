# -*- coding: utf-8 -*-
r"""Gera ficheiros .dotx (modelo Word) para preenchimento e importacao assistida SRA.
Executar na raiz do repositorio (PowerShell)::

  .\.venv\Scripts\python.exe scripts/build_canonical_upload_dotx.py

Saida: modelos_upload_doc_canonicos/
"""
from __future__ import annotations

import io
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from app.models import SECOES_PADRAO  # noqa: E402
from app.notificacoes.modelos import filename_para, slug_titulo  # noqa: E402

NUM_TO_TITLE: dict[str, str] = dict(SECOES_PADRAO)

OUT_DIR = ROOT / "modelos_upload_doc_canonicos"
NAVY = RGBColor(0x1C, 0x3D, 0x59)
MUTED = RGBColor(0x4F, 0x5D, 0x6E)

_DOC_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
_TPL_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml"


def _docx_bytes_to_dotx(data: bytes) -> bytes:
    """Converte um .docx em .dotx ajustando o content type de documento principal."""
    in_buf = io.BytesIO(data)
    out_buf = io.BytesIO()
    with zipfile.ZipFile(in_buf, "r") as zin:
        with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                raw = zin.read(item)
                if item == "[Content_Types].xml":
                    txt = raw.decode("utf-8")
                    txt = txt.replace(_DOC_CT, _TPL_CT)
                    raw = txt.encode("utf-8")
                zout.writestr(item, raw)
    return out_buf.getvalue()


# Slug e nome de ficheiro vêm de `app.notificacoes.modelos` (fonte única).
# Mantém alias local para compatibilidade com chamadas internas existentes.
_slug_title = slug_titulo


def _heading_level(numero: str) -> int:
    dots = (numero or "").count(".")
    return min(1 + dots, 3)


def _section_chain(numero: str) -> list[tuple[str, str]]:
    """Cadeia numerica presente em SECOES_PADRAO ate `numero` (ex.: 4.4.7 -> 4, 4.4, 4.4.7)."""
    raw = str(numero or "").strip()
    if not raw:
        return []
    parts = raw.split(".")
    chain = []
    for i in range(len(parts)):
        prefix = ".".join(parts[: i + 1])
        if prefix in NUM_TO_TITLE:
            chain.append((prefix, NUM_TO_TITLE[prefix]))
    return chain


def _apply_base_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Verdana"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0x14, 0x14, 0x14)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for lvl in (1, 2, 3):
        try:
            h = document.styles[f"Heading {lvl}"]
        except KeyError:
            continue
        h.font.name = "Verdana"
        h.font.bold = True
        h.font.color.rgb = NAVY
        h.font.size = Pt(14 if lvl == 1 else 12 if lvl == 2 else 11)
    try:
        lp = document.styles["List Paragraph"]
        lp.font.name = "Verdana"
        lp.font.size = Pt(10)
    except KeyError:
        pass


def _fmt_caption(p) -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def _add_instruction_block(document: Document) -> None:
    t = document.add_paragraph()
    t.add_run("Modelo SRA (importacao assistida). Nao apague as linhas de exemplo de ").bold = False
    r = t.add_run("legenda e marcadores de lista")
    r.bold = True
    t.add_run(
        " — pode substituir o texto, mas mantenha o estilo. Copie, cole e mova blocos dentro desta secao."
    )
    t.paragraph_format.space_after = Pt(12)
    for run in t.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = MUTED


def _add_bloco_texto(document: Document) -> None:
    document.add_paragraph("---", style="Normal")
    h = document.add_paragraph("Bloco: paragrafo (texto)")
    h.runs[0].font.bold = True
    h.runs[0].font.size = Pt(10)
    h.runs[0].font.color.rgb = NAVY
    document.add_paragraph(
        "Escreva aqui textos corridos, subtitulos no padrao de numeracao (ex.: 4.4.7.1 Titulo do "
        + "subtema) e linhas vazias para separar paragrafos."
    )


def _add_bloco_lista(document: Document) -> None:
    document.add_paragraph("---", style="Normal")
    h = document.add_paragraph("Bloco: listas (nao e numeracao de secoes do relatorio)")
    h.runs[0].font.bold = True
    h.runs[0].font.size = Pt(10)
    h.runs[0].font.color.rgb = NAVY
    p1 = document.add_paragraph("Item com marcador (lista real do Word, nivel 1).", style="List Bullet")
    p1.paragraph_format.left_indent = Cm(0.5)
    p1.paragraph_format.first_line_indent = Cm(-0.25)
    p2 = document.add_paragraph("Segundo item (mesmo estilo, nivel 1).", style="List Bullet")
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.first_line_indent = Cm(-0.25)
    p3 = document.add_paragraph("Item numerado (lista real com numeros).", style="List Number")
    p3.paragraph_format.left_indent = Cm(0.5)
    p3.paragraph_format.first_line_indent = Cm(-0.25)
    p4 = document.add_paragraph("Segundo item numerado.", style="List Number")
    p4.paragraph_format.left_indent = Cm(0.5)
    p4.paragraph_format.first_line_indent = Cm(-0.25)


def _add_bloco_figura(document: Document) -> None:
    document.add_paragraph("---", style="Normal")
    h = document.add_paragraph("Bloco: figura (imagem + legenda + fonte)")
    h.runs[0].font.bold = True
    h.runs[0].font.size = Pt(10)
    h.runs[0].font.color.rgb = NAVY
    p = document.add_paragraph(
        "CLIQUE AQUI, depois insera a figura: Inserir > Imagens. Deixe a legenda e a fonte nas linhas "
        + "imediatas a seguir (formato abaixo)."
    )
    for run in p.runs:
        run.font.italic = True
        run.font.size = Pt(9)
    cap = document.add_paragraph("Figura X.Y: descricao clara do que a imagem mostra.")
    _fmt_caption(cap)
    f = document.add_paragraph("Fonte: Nome e ano da fonte ou referencia.")
    _fmt_caption(f)


def _add_bloco_tabela(document: Document) -> None:
    document.add_paragraph("---", style="Normal")
    h = document.add_paragraph("Bloco: tabela (grelha Word + legenda + fonte)")
    h.runs[0].font.bold = True
    h.runs[0].font.size = Pt(10)
    h.runs[0].font.color.rgb = NAVY
    tcap = document.add_paragraph("Tabela X.Y: legenda com contexto (substitua X.Y se quiser, o sistema ajusta).")
    tcap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in tcap.runs:
        run.font.italic = True
        run.font.size = Pt(9)
    table = document.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    for ri, label in enumerate(("Col A", "Col B", "Col C")):
        table.rows[0].cells[ri].text = label
    table.rows[1].cells[0].text = "Celula"
    table.rows[1].cells[1].text = "Celula"
    table.rows[1].cells[2].text = "Celula"
    table.rows[2].cells[0].text = "Celula"
    table.rows[2].cells[1].text = "Celula"
    table.rows[2].cells[2].text = "Celula"
    tf = document.add_paragraph("Fonte: origem dos dados (instituicao, planilha, data).")
    _fmt_caption(tf)


def _add_section_header(document: Document, num: str, tit: str) -> None:
    lv = _heading_level(num)
    p = document.add_heading(f"{num}  {tit}", level=lv)
    for r in p.runs:
        r.font.name = "Verdana"
        r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(8)


def _build_document_for_secao(
    secoes: list[tuple[str, str]], *, all_sections: bool
) -> Document:
    document = Document()
    _apply_base_styles(document)
    t = document.add_paragraph("Relatorio de atividades (rascunho para importar no SRA)")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].font.size = Pt(12)
    t.runs[0].font.bold = True
    t.runs[0].font.color.rgb = NAVY
    t.paragraph_format.space_after = Pt(6)
    sub = document.add_paragraph(
        "PLI/SP-2050 — conteudo da secao indicada, sem alterar a estrutura deste modelo (apenas conteudo)."
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
    sub.paragraph_format.space_after = Pt(12)
    _add_instruction_block(document)
    for i, (num, tit) in enumerate(secoes):
        if all_sections and i:
            document.add_page_break()
        if not all_sections and len(secoes) == 1:
            chain = _section_chain(num)
            if chain and len(chain) > 1:
                for cn, ctit in chain[:-1]:
                    _add_section_header(document, cn, ctit)
                    ctx = document.add_paragraph(
                        "Titulo ancestral (somente contexto neste modelo). Escreva o texto na ultima "
                        "subseccao deste documento ou no ficheiro .dotx dessa parte."
                    )
                    for ru in ctx.runs:
                        ru.font.size = Pt(9)
                        ru.font.italic = True
                        ru.font.color.rgb = MUTED
                    ctx.paragraph_format.space_after = Pt(12)
                _add_section_header(document, chain[-1][0], chain[-1][1])
                _add_bloco_texto(document)
                _add_bloco_lista(document)
                _add_bloco_figura(document)
                _add_bloco_tabela(document)
                continue

        _add_section_header(document, num, tit)
        if not all_sections:
            _add_bloco_texto(document)
            _add_bloco_lista(document)
            _add_bloco_figura(document)
            _add_bloco_tabela(document)
        else:
            p = document.add_paragraph(
                "Reserva de espaco: preencha o conteudo desta subseccao. Para exemplos de lista, "
                + "figura e tabela, abra o ficheiro .dotx unico desta seccao (ex.: secao_4_4_7_...)."
            )
            for run in p.runs:
                run.font.size = Pt(9)
    return document


def _section_filename(num: str, tit: str) -> str:
    return filename_para(num, tit)


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    # Completo: todas as secoes, uma pagina de instrucao + uma por secao (layout curto)
    full_doc = _build_document_for_secao(list(SECOES_PADRAO), all_sections=True)
    buf = io.BytesIO()
    full_doc.save(buf)
    full_dotx = _docx_bytes_to_dotx(buf.getvalue())
    (OUT_DIR / "SRA_todas_secoes.dotx").write_bytes(full_dotx)
    for num, tit in SECOES_PADRAO:
        doc = _build_document_for_secao([(num, tit)], all_sections=False)
        b = io.BytesIO()
        doc.save(b)
        name = _section_filename(num, tit)
        (OUT_DIR / name).write_bytes(_docx_bytes_to_dotx(b.getvalue()))
    readme = (
        "Modelos Word (.dotx) para preenchimento e importacao da secao no SRA (importacao assistida).\n"
        + "- `SRA_todas_secoes.dotx`: grelha com titulos de todas as secoes padrao (referencia do sumario).\n"
        + "- `secao_*.dotx`: um ficheiro por secao — use o que corresponder ao seu numero (ex. 4.4.7…). Se o "
        + "numero tiver ascendencia no sumario, o modelo inclui ate tres titulos Word (Heading 1 a 3) encadeados; "
        + "somente a ultima subseccao traz os exemplos texto/lista/figura/tabela.\n"
        + "Preencha apenas conteudo; prefira listas, tabelas e legendas reais de Word, conforme as "
        + "caixas de exemplo. Guarde o ficheiro e envie em Importar na edicao de secao.\n"
    )
    (OUT_DIR / "README.md").write_text(readme, encoding="utf-8")
    n = 1 + len(SECOES_PADRAO)
    print(f"Escritos {n} ficheiros .dotx em {OUT_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
